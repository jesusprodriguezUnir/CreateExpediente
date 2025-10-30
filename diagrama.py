# -*- coding: utf-8 -*-
"""Módulo para generar el diagrama y documento Word con los flujos.

Funciones públicas:
- generate_diagram(output_dir=None): genera PNG y DOCX en output_dir y devuelve (img_path, doc_path)
- main(): ejecuta generate_diagram() usando el directorio de salida por defecto
"""
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches
import os


def _load_fonts():
    """Carga las fuentes con mejor soporte para caracteres especiales y acentos."""
    try:
        # Intentar fuentes del sistema con buen soporte Unicode
        font_candidates = [
            # Windows fonts
            "arial.ttf", "calibri.ttf", "segoeui.ttf", "tahoma.ttf",
            # Cross-platform fonts
            "DejaVuSans-Bold.ttf", "liberation-sans-bold.ttf",
            # Fallback fonts
            "NotoSans-Bold.ttf", "OpenSans-Bold.ttf"
        ]
        
        font_title = None
        font_box = None
        font_small = None
        
        # Buscar fuentes disponibles
        for font_name in font_candidates:
            try:
                font_title = ImageFont.truetype(font_name, 26)  # Reducir de 28 a 26
                font_box = ImageFont.truetype(font_name.replace("-Bold", ""), 16)  # Reducir de 18 a 16
                font_small = ImageFont.truetype(font_name.replace("-Bold", ""), 12)  # Reducir de 14 a 12
                print(f"Usando fuente: {font_name}")
                break
            except (OSError, IOError):
                continue
        
        # Si no se encuentra ninguna fuente específica, usar fuentes del sistema
        if font_title is None:
            try:
                # Intentar fuentes del sistema Windows
                font_title = ImageFont.truetype("arial.ttf", 26)  # Reducir tamaño
                font_box = ImageFont.truetype("arial.ttf", 16)    # Reducir tamaño
                font_small = ImageFont.truetype("arial.ttf", 12)  # Reducir tamaño
                print("Usando fuente del sistema: Arial")
            except (OSError, IOError):
                # Último recurso: fuente por defecto de PIL (pero mejorada)
                font_title = ImageFont.load_default()
                font_box = ImageFont.load_default()
                font_small = ImageFont.load_default()
                print("Usando fuente por defecto de PIL")
                
    except Exception as e:
        print(f"Error cargando fuentes: {e}")
        # Fallback seguro
        font_title = ImageFont.load_default()
        font_box = ImageFont.load_default()  
        font_small = ImageFont.load_default()
        
    return font_title, font_box, font_small


def _generate_png_diagram(img_path):
    """Genera únicamente la imagen PNG del diagrama.
    
    Args:
        img_path (str): ruta completa donde guardar la imagen PNG
        
    Returns:
        str: ruta de la imagen generada
    """
    # Create an image with PIL for clean boxes and arrows
    width, height = 2000, 970  # Aumentar ancho para cajas verdes movidas hacia la derecha
    bg_color = (248, 250, 252)  # Fondo gris muy claro para mejor contraste
    img = Image.new("RGB", (width, height), bg_color)
    draw = ImageDraw.Draw(img)

    font_title, font_box, font_small = _load_fonts()

    # Colors (user specified) - Mejorados para mejor contraste
    color_gestor = (14, 82, 160)       # dark blue for GestorMapeos
    color_erp = (54, 126, 223)         # medium blue for ERP Académico
    color_expedientes = (76, 175, 80)  # green for Expedientes - más vibrante
    border_color = (30, 50, 80)
    shadow_color = (200, 200, 200)     # Color para sombras

    # Positions for boxes - Mejor distribución y espaciado con cajas ajustadas
    x_gestor, y_gestor = 120, 280
    w_box_gestor, h_box_gestor = 350, 100  # Reducir GestorMapeos de 380x110 a 350x100
    w_box, h_box = 470, 150  # Tamaño para cajas azules (ERP)
    w_box_green, h_box_green = 490, 160  # Cajas verdes un pelín más grandes

    x_api, y_api = 600, 150
    x_api_ampliacion, y_api_ampliacion = 600, 540  # Ajustar por cajas más altas

    # Three green boxes stacked vertically on the right with better spacing
    x_expe_post1, y_expe_post1 = 1200, 150  # Mover hacia la derecha de 1150 a 1200
    x_expe_post2, y_expe_post2 = 1200, 330  # POST /api/v1/expedientes-alumnos (más separación)
    x_expe_post3, y_expe_post3 = 1200, 510  # POST /api/v1/expedientes-alumnos/matricula-realizada

    # Draw title with better styling
    title_text = "Flujos: GestorMapeos - ERP Académico - Expedientes"
    # Dibujar sombra del título
    draw.text((width // 2 - 218, 22), title_text, font=font_title, fill=(180,180,180))
    # Dibujar título principal
    draw.text((width // 2 - 220, 20), title_text, font=font_title, fill=(20,20,20))
    
    # Agregar línea decorativa bajo el título
    draw.line((width // 2 - 350, 70, width // 2 + 350, 70), fill=(100,100,100), width=2)

    # Helper to draw rounded rectangle with text and shadow
    def draw_box(x, y, w, h, fill, text, font, outline=(30,50,80)):
        radius = 15
        # Dibujar sombra
        shadow_offset = 4
        draw.rounded_rectangle([x+shadow_offset, y+shadow_offset, x+w+shadow_offset, y+h+shadow_offset], 
                              radius=radius, fill=shadow_color, outline=None)
        # outer rect
        draw.rounded_rectangle([x, y, x+w, y+h], radius=radius, fill=fill, outline=outline, width=3)
        # center text (multilínea)
        lines = text.split("\n")
        line_sizes = []
        for line in lines:
            if line.strip() == "":
                bbox = draw.textbbox((0, 0), "A", font=font)
            else:
                bbox = draw.textbbox((0, 0), line, font=font)
            tw = bbox[2] - bbox[0]
            th = bbox[3] - bbox[1]
            line_sizes.append((tw, th))

        total_h = sum(hgt for _, hgt in line_sizes)
        start_y = y + (h - total_h) / 2
        y_offset = 0
        for (line, (tw, th)) in zip(lines, line_sizes):
            draw.text((x + (w - tw) / 2, start_y + y_offset), line, font=font,
                      fill=(255,255,255) if sum(fill[:3]) < 400 else (0,0,0))
            y_offset += th

    # Draw boxes
    draw_box(x_gestor, y_gestor, w_box_gestor, h_box_gestor, color_gestor, "GestorMapeos", font_box)
    draw_box(x_api, y_api, w_box, h_box, color_erp,
             "API Primera Matrícula\n\nPOST https://erpacademico.unir.net/api/v1/migrar\n\n",
             font_box)
    draw_box(x_api_ampliacion, y_api_ampliacion, w_box, h_box, color_erp,
             "API Ampliación\n\nPOST https://erpacademico.unir.net/api/v1/migrar/ampliacion\n\n",
             font_box)
    # Three green boxes stacked vertically with larger size
    draw_box(x_expe_post1, y_expe_post1, w_box_green, h_box_green, color_expedientes, 
             "Expedientes\n\nhttps://expedientesacademico.unir.net", font_box)
    draw_box(x_expe_post2, y_expe_post2, w_box_green, h_box_green, color_expedientes, 
             "POST /api/v1/expedientes-alumnos", font_box)
    draw_box(x_expe_post3, y_expe_post3, w_box_green, h_box_green, color_expedientes, 
             "POST /api/v1/expedientes-alumnos/matricula-realizada", font_box)

    # Draw arrows (lines with triangles) - Improved
    def draw_arrow(x1, y1, x2, y2, fill=(40,40,40), width_line=5):
        draw.line((x1, y1, x2, y2), fill=fill, width=width_line)
        # draw triangle head
        import math
        angle = math.atan2(y2 - y1, x2 - x1)
        head_len = 16
        left = (x2 - head_len * math.cos(angle) + head_len/2 * math.sin(angle),
                y2 - head_len * math.sin(angle) - head_len/2 * math.cos(angle))
        right = (x2 - head_len * math.cos(angle) - head_len/2 * math.sin(angle),
                 y2 - head_len * math.sin(angle) + head_len/2 * math.cos(angle))
        draw.polygon([ (x2,y2), left, right ], fill=fill)

    # Function to draw text on arrows with better background
    def draw_arrow_with_text(x1, y1, x2, y2, text, fill=(40,40,40), width_line=5):
        draw_arrow(x1, y1, x2, y2, fill, width_line)
        # Calculate text position (middle of the arrow)
        text_x = (x1 + x2) // 2 - 50  # Ajustar para cajas más grandes
        text_y = (y1 + y2) // 2 - 20  # Mantener posición vertical
        # Fondo blanco con borde para el texto
        bbox = draw.textbbox((text_x, text_y), text, font=font_small)
        padding = 6  # Más padding para texto más pequeño
        draw.rectangle([bbox[0]-padding, bbox[1]-padding, bbox[2]+padding, bbox[3]+padding], 
                      fill=(255,255,255), outline=(150,150,150), width=2)
        draw.text((text_x, text_y), text, font=font_small, fill=fill)

    # Flow 1: GestorMapeos -> API Primera Matrícula -> Two green boxes (not matricula-realizada)
    draw_arrow_with_text(x_gestor + w_box_gestor, y_gestor + h_box_gestor/2, 
                         x_api, y_api + h_box/2, 
                         "Primera Matrícula")
    
    # From API Primera Matrícula to only the first two green boxes
    draw_arrow_with_text(x_api + w_box, y_api + h_box/2, 
                         x_expe_post1, y_expe_post1 + h_box_green/2, 
                         "Actualizar Expediente")
    
    draw_arrow_with_text(x_api + w_box, y_api + h_box/2 + 10, 
                         x_expe_post2, y_expe_post2 + h_box_green/2, 
                         "Crear Expediente")

    # Flow 2: GestorMapeos -> API Ampliación -> Last green box only
    draw_arrow_with_text(x_gestor + w_box_gestor, y_gestor + h_box_gestor/2 + 40, 
                         x_api_ampliacion, y_api_ampliacion + h_box/2, 
                         "Ampliación")
    
    draw_arrow_with_text(x_api_ampliacion + w_box, y_api_ampliacion + h_box/2, 
                         x_expe_post3, y_expe_post3 + h_box_green/2, 
                         "Matrícula Realizada")

    # Add legends at bottom with better styling
    legend_y = 820  # Ajustar para cajas más altas
    legend_box_size = 28
    
    # Fondo para la leyenda
    draw.rectangle([60, legend_y-15, 720, legend_y+50], fill=(255,255,255), outline=(180,180,180), width=2)
    
    draw.rectangle([80, legend_y, 80+legend_box_size, legend_y+legend_box_size], 
                   fill=color_gestor, outline=border_color, width=2)
    draw.text((120, legend_y+2), "GestorMapeos", font=font_small, fill=(0,0,0))
    
    draw.rectangle([280, legend_y, 280+legend_box_size, legend_y+legend_box_size], 
                   fill=color_erp, outline=border_color, width=2)
    draw.text((320, legend_y+2), "ERP Académico", font=font_small, fill=(0,0,0))
    
    draw.rectangle([480, legend_y, 480+legend_box_size, legend_y+legend_box_size], 
                   fill=color_expedientes, outline=border_color, width=2)
    draw.text((520, legend_y+2), "Expedientes", font=font_small, fill=(0,0,0))

    # Save image
    img.save(img_path)
    return img_path


def _generate_word_document(doc_path, img_path):
    """Genera únicamente el documento Word con documentación detallada.
    
    Args:
        doc_path (str): ruta completa donde guardar el documento DOCX
        img_path (str): ruta de la imagen PNG para embeber en el documento
        
    Returns:
        str: ruta del documento generado
    """
    # Create Word doc and embed image and textual info
    doc = Document()
    
    # Título principal
    doc.add_heading('Esquema de Flujos: GestorMapeos - ERP Académico - Expedientes', level=1)
    
    # Introducción
    intro = doc.add_paragraph()
    intro.add_run('Este documento describe los flujos de integración entre los sistemas GestorMapeos, ERP Académico y Expedientes. ').bold = True
    doc.add_paragraph('El diagrama visual muestra los dos flujos principales: "Primera Matrícula" y "Ampliación", con colores diferenciados por sistema y URLs completas de los endpoints utilizados.')
    
    # Imagen del diagrama al inicio
    doc.add_heading('Diagrama Visual', level=2)
    doc.add_picture(img_path, width=Inches(6.5))
    doc.add_paragraph()  # Espacio

    doc.add_heading('Descripción de Sistemas', level=2)
    systems_para = doc.add_paragraph()
    systems_para.add_run('• GestorMapeos: ').bold = True
    systems_para.add_run('Sistema de gestión de mapeos de matrículas (azul oscuro)\n')
    systems_para.add_run('• ERP Académico: ').bold = True
    systems_para.add_run('Sistema de planificación de recursos académicos (azul medio)\n')
    systems_para.add_run('• Expedientes: ').bold = True
    systems_para.add_run('Sistema de gestión de expedientes académicos (verde)')

    doc.add_heading('Endpoints de las APIs', level=2)
    
    # ERP Académico
    doc.add_heading('ERP Académico', level=3)
    doc.add_paragraph('• Primera migración (crear/actualizar expediente):', style='List Bullet')
    doc.add_paragraph('  POST https://erpacademico.unir.net/api/v1/migrar', style='List Bullet 2')
    doc.add_paragraph('• Ampliación (no primera matrícula):', style='List Bullet')
    doc.add_paragraph('  POST https://erpacademico.unir.net/api/v1/migrar/ampliacion', style='List Bullet 2')
    
    # Expedientes
    doc.add_heading('Expedientes', level=3)
    doc.add_paragraph('• Crear expediente:', style='List Bullet')
    doc.add_paragraph('  POST https://expedienteserp.unir.net/api/v1/expedientes-alumnos', style='List Bullet 2')
    doc.add_paragraph('• Modificar expediente por integración:', style='List Bullet')
    doc.add_paragraph('  PUT https://expedienteserp.unir.net/api/v1/expedientes-alumnos/{id}/por-integracion', style='List Bullet 2')
    doc.add_paragraph('• Notificar matrícula realizada:', style='List Bullet')
    doc.add_paragraph('  POST https://expedienteserp.unir.net/api/v1/expedientes-alumnos/matricula-realizada', style='List Bullet 2')

    doc.add_heading('Descripción Detallada de los Flujos', level=2)
    
    # Flujo 1
    doc.add_heading('Flujo 1 — Primera Matrícula', level=3)
    flow1_para = doc.add_paragraph()
    flow1_para.add_run('Proceso:\n').bold = True
    doc.add_paragraph('1. GestorMapeos envía la información de la primera matrícula', style='List Number')
    doc.add_paragraph('   → POST https://erpacademico.unir.net/api/v1/migrar', style='List Number 2')
    doc.add_paragraph('2. El ERP Académico procesa y guarda la matrícula', style='List Number')
    doc.add_paragraph('3. Si es la PRIMERA matrícula, el ERP puede:', style='List Number')
    doc.add_paragraph('   a) Crear un nuevo expediente en Expedientes:', style='List Number 2')
    doc.add_paragraph('      POST https://expedienteserpunir.net/api/v1/expedientes-alumnos', style='List Number 3')
    doc.add_paragraph('   b) Actualizar un expediente existente:', style='List Number 2')
    doc.add_paragraph('      PUT https://expedienteserp.unir.net/api/v1/expedientes-alumnos/{id}/por-integracion', style='List Number 3')

    # Flujo 2
    doc.add_heading('Flujo 2 — Ampliación (No Primera Matrícula)', level=3)
    flow2_para = doc.add_paragraph()
    flow2_para.add_run('Proceso:\n').bold = True
    doc.add_paragraph('1. GestorMapeos envía la información de ampliación', style='List Number')
    doc.add_paragraph('   → POST https://erpacademico.unir.net/api/v1/migrar/ampliacion', style='List Number 2')
    doc.add_paragraph('2. El ERP Académico procesa y guarda la matrícula', style='List Number')
    doc.add_paragraph('3. El ERP notifica a Expedientes que se ha realizado una matrícula:', style='List Number')
    doc.add_paragraph('   → POST https://expedienteserp.unir.net/api/v1/expedientes-alumnos/matricula-realizada', style='List Number 2')

    doc.add_heading('Código PlantUML', level=2)
    plantuml_code = """@startuml
skinparam rectangle {
  BackgroundColor<<Gestor>> #0E52A0
  BackgroundColor<<ERP>> #367EDF
  BackgroundColor<<Expedientes>> #4CAF50
  FontColor white
}
actor "GestorMapeos" as GM <<Gestor>>
rectangle "ERP Académico\nPOST /api/v1/migrar (primera)\nPOST /api/v1/migrar/ampliacion (no primera)" as ERP <<ERP>>
rectangle "Expedientes\nhttps://expedienteserp.unir.net" as EXP <<Expedientes>>
GM --> ERP : POST https://erpacademico.unir.net/api/v1/migrar
ERP --> EXP : "Primera matrícula -> crear/actualizar expediente"
note right of ERP
    - Crear: POST https://expedienteserp.unir.net/api/v1/expedientes-alumnos
    - Actualizar: PUT https://expedienteserp.unir.net/api/v1/expedientes-alumnos/{id}/por-integracion
end note
GM --> ERP : POST https://erpacademico.unir.net/api/v1/migrar/ampliacion
ERP --> EXP : POST https://expedienteserp.unir.net/api/v1/expedientes-alumnos/matricula-realizada
@enduml"""
    doc.add_paragraph(plantuml_code)
    
    # Nota final
    doc.add_paragraph()
    note_para = doc.add_paragraph()
    note_para.add_run('Nota: ').bold = True
    note_para.add_run('Este diagrama puede ser renderizado usando PlantUML para generar diagramas UML alternativos.')

    doc.save(doc_path)
    return doc_path


def generate_png_only(output_dir=None):
    """Genera únicamente la imagen PNG del diagrama.

    Args:
        output_dir (str): carpeta donde guardar el resultado. Si es None, se usa
            el subdirectorio `output` dentro del paquete.

    Returns:
        str: ruta de la imagen PNG generada
    """
    pkg_dir = os.path.dirname(__file__)
    if output_dir is None:
        output_dir = os.path.join(pkg_dir, "output")
    os.makedirs(output_dir, exist_ok=True)

    img_path = os.path.join(output_dir, "diagram_expedientes_flow.png")
    return _generate_png_diagram(img_path)


def generate_word_only(output_dir=None, img_path=None):
    """Genera únicamente el documento Word.

    Args:
        output_dir (str): carpeta donde guardar el resultado. Si es None, se usa
            el subdirectorio `output` dentro del paquete.
        img_path (str): ruta de la imagen PNG para embeber. Si es None, se busca
            en el directorio de salida.

    Returns:
        str: ruta del documento DOCX generado
    """
    pkg_dir = os.path.dirname(__file__)
    if output_dir is None:
        output_dir = os.path.join(pkg_dir, "output")
    os.makedirs(output_dir, exist_ok=True)

    doc_path = os.path.join(output_dir, "Esquema_Flujos_GestorMapeos_ERP_Expedientes.docx")
    
    if img_path is None:
        img_path = os.path.join(output_dir, "diagram_expedientes_flow.png")
        # Si no existe la imagen, crearla primero
        if not os.path.exists(img_path):
            _generate_png_diagram(img_path)
    
    return _generate_word_document(doc_path, img_path)


def generate_diagram(output_dir=None):
    """Genera la imagen PNG y el documento DOCX.

    Args:
        output_dir (str): carpeta donde guardar los resultados. Si es None, se usa
            el subdirectorio `output` dentro del paquete.

    Returns:
        tuple: (img_path, doc_path)
    """
    pkg_dir = os.path.dirname(__file__)
    if output_dir is None:
        output_dir = os.path.join(pkg_dir, "output")
    os.makedirs(output_dir, exist_ok=True)

    img_path = os.path.join(output_dir, "diagram_expedientes_flow.png")
    doc_path = os.path.join(output_dir, "Esquema_Flujos_GestorMapeos_ERP_Expedientes.docx")

    # Generar la imagen PNG
    img_path = _generate_png_diagram(img_path)
    
    # Generar el documento Word
    doc_path = _generate_word_document(doc_path, img_path)

    return img_path, doc_path


def main():
    img, doc = generate_diagram()
    print(f"image:{img}")
    print(f"doc:{doc}")


if __name__ == '__main__':
    main()
